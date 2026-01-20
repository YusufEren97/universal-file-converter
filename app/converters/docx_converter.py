"""
DOCX Converter
Converts DOCX files to PDF, TXT, HTML, MD formats.
Preserves formatting when converting to PDF.
"""
import os
import asyncio
import subprocess
import shutil


async def convert_docx(input_path: str, output_dir: str, target_format: str) -> dict:
    """DOCX converter supporting PDF, TXT, HTML, MD outputs."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _process_docx, input_path, output_dir, target_format)


def _process_docx(input_path: str, output_dir: str, target_format: str) -> dict:
    try:
        filename = os.path.basename(input_path)
        name, ext = os.path.splitext(filename)
        output_filename = f"{name}.{target_format}"
        output_path = os.path.join(output_dir, output_filename)

        if target_format == 'pdf':
            return _docx_to_pdf(input_path, output_path, output_filename)
        elif target_format == 'txt':
            return _docx_to_txt(input_path, output_path, output_filename)
        elif target_format == 'html':
            return _docx_to_html(input_path, output_path, output_filename, name)
        elif target_format == 'md':
            return _docx_to_md(input_path, output_path, output_filename, name)
        else:
            return {"success": False, "error": f"Unsupported target format for DOCX: {target_format}"}

    except Exception as e:
        return {"success": False, "error": f"DOCX conversion error: {str(e)}"}


def _docx_to_pdf(input_path: str, output_path: str, output_filename: str) -> dict:
    """Convert DOCX to PDF using docx2pdf (requires Word) or LibreOffice."""
    
    # Method 1: Try docx2pdf (Windows + Microsoft Word)
    try:
        from docx2pdf import convert
        convert(input_path, output_path)
        if os.path.exists(output_path):
            return {"success": True, "output_path": output_path, "filename": output_filename}
    except Exception as e:
        print(f"[DOCX→PDF] docx2pdf failed: {e}")
    
    # Method 2: Try LibreOffice
    try:
        libreoffice_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice",  # If in PATH
            "/usr/bin/soffice",  # Linux
            "/Applications/LibreOffice.app/Contents/MacOS/soffice"  # macOS
        ]
        
        soffice_path = None
        for path in libreoffice_paths:
            if os.path.exists(path) or shutil.which(path):
                soffice_path = path
                break
        
        if soffice_path:
            output_dir = os.path.dirname(output_path)
            result = subprocess.run([
                soffice_path,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                input_path
            ], capture_output=True, text=True, timeout=120)
            
            # LibreOffice outputs with original name
            expected_output = os.path.join(output_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
            if os.path.exists(expected_output):
                if expected_output != output_path:
                    shutil.move(expected_output, output_path)
                return {"success": True, "output_path": output_path, "filename": output_filename}
    except Exception as e:
        print(f"[DOCX→PDF] LibreOffice failed: {e}")
    
    return {
        "success": False, 
        "error": "PDF conversion requires Microsoft Word or LibreOffice. Please install one of them."
    }


def _docx_to_txt(input_path: str, output_path: str, output_filename: str) -> dict:
    """Extract text from DOCX."""
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run 'pip install python-docx'"}
    
    doc = Document(input_path)
    text_content = []
    
    for para in doc.paragraphs:
        text_content.append(para.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text_content.append("\t".join(row_text))
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(text_content))
    
    return {"success": True, "output_path": output_path, "filename": output_filename}


def _docx_to_html(input_path: str, output_path: str, output_filename: str, name: str) -> dict:
    """Convert DOCX to HTML with basic styling."""
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run 'pip install python-docx'"}
    
    doc = Document(input_path)
    
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{name}</title>
    <style>
        body {{ 
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; 
            max-width: 800px; 
            margin: 40px auto; 
            padding: 20px; 
            line-height: 1.6; 
        }}
        h1 {{ font-size: 2em; margin-top: 1em; }}
        h2 {{ font-size: 1.5em; margin-top: 0.8em; }}
        h3 {{ font-size: 1.2em; margin-top: 0.6em; }}
        p {{ margin: 1em 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f5f5f5; }}
    </style>
</head>
<body>
"""
    
    for para in doc.paragraphs:
        if para.text.strip():
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                html_content += f"    <h1>{para.text}</h1>\n"
            elif "Heading 2" in style_name:
                html_content += f"    <h2>{para.text}</h2>\n"
            elif "Heading 3" in style_name:
                html_content += f"    <h3>{para.text}</h3>\n"
            else:
                html_content += f"    <p>{para.text}</p>\n"
    
    # Convert tables
    for table in doc.tables:
        html_content += "    <table>\n"
        for i, row in enumerate(table.rows):
            html_content += "        <tr>\n"
            tag = "th" if i == 0 else "td"
            for cell in row.cells:
                html_content += f"            <{tag}>{cell.text}</{tag}>\n"
            html_content += "        </tr>\n"
        html_content += "    </table>\n"
    
    html_content += "</body>\n</html>"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return {"success": True, "output_path": output_path, "filename": output_filename}


def _docx_to_md(input_path: str, output_path: str, output_filename: str, name: str) -> dict:
    """Convert DOCX to Markdown."""
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed. Run 'pip install python-docx'"}
    
    doc = Document(input_path)
    md_content = f"# {name}\n\n"
    
    for para in doc.paragraphs:
        if para.text.strip():
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                md_content += f"# {para.text}\n\n"
            elif "Heading 2" in style_name:
                md_content += f"## {para.text}\n\n"
            elif "Heading 3" in style_name:
                md_content += f"### {para.text}\n\n"
            else:
                md_content += f"{para.text}\n\n"
    
    # Convert tables to markdown
    for table in doc.tables:
        if table.rows:
            # Header row
            header = [cell.text for cell in table.rows[0].cells]
            md_content += "| " + " | ".join(header) + " |\n"
            md_content += "| " + " | ".join(["---"] * len(header)) + " |\n"
            
            # Data rows
            for row in table.rows[1:]:
                cells = [cell.text for cell in row.cells]
                md_content += "| " + " | ".join(cells) + " |\n"
            md_content += "\n"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    
    return {"success": True, "output_path": output_path, "filename": output_filename}
